"""Word (.docx) parser。

策略：
- 用 python-docx 读段落 + 表格
- 标题段（Heading 1/2/3/...）成为 TocEntry
- 所有正文段落 → TextBlock（按文档顺序 reading_order）
- 表格 → ParsedTable（headers + rows）
- 图片暂不抽（M3 时机不成熟，python-docx 提取嵌入图比较繁琐，留待 M4）

对 docx 的章节路径推断：
- Heading 1 → "1"
- 下一个 Heading 2 → "1.1"
- 维护栈式计数器，与 SectionExtractor 配合

为避免硬依赖，python-docx 按需 import。
"""
from __future__ import annotations

from pathlib import Path

from docgraph.graph.schema import (
    ParsedDoc,
    ParsedPage,
    ParsedTable,
    TextBlock,
    TocEntry,
)
from docgraph.parsers.base import ParseContext
from docgraph.parsers.normalize import populate_l0_blocks


class DocxParser:
    name = "docx"
    supports = {".docx"}
    version = "0.1"

    def can_parse(self, path: Path) -> bool:
        return path.suffix.lower() in self.supports

    def parse(self, path: Path, ctx: ParseContext) -> ParsedDoc:
        try:
            from docx import Document  # type: ignore  # python-docx
        except ImportError as e:  # pragma: no cover
            raise RuntimeError(
                "python-docx not installed. Install with: "
                "pip install 'docgraph[documents]'"
            ) from e

        doc = Document(str(path))

        # 1. 提取所有段落 + 标题
        toc: list[TocEntry] = []
        text_blocks: list[TextBlock] = []
        heading_counter = [0, 0, 0, 0, 0, 0]  # 1..6 级
        order = 0

        for para in doc.paragraphs:
            text = (para.text or "").strip()
            if not text:
                continue
            style = (para.style.name or "") if para.style else ""
            heading_level = self._heading_level(style)
            if heading_level:
                # 维护章节号
                heading_counter[heading_level - 1] += 1
                for i in range(heading_level, 6):
                    heading_counter[i] = 0
                path_str = ".".join(
                    str(n) for n in heading_counter[:heading_level] if n > 0
                )
                toc.append(
                    TocEntry(
                        level=heading_level,
                        title=text,
                        page=None,
                        section_path=path_str or None,
                    )
                )
            text_blocks.append(
                TextBlock(
                    text=text,
                    reading_order=order,
                    is_heading=bool(heading_level),
                    heading_level=heading_level,
                )
            )
            order += 1

        # 2. 提取表格
        tables: list[ParsedTable] = []
        for tbl in doc.tables:
            rows: list[list[str]] = []
            for row in tbl.rows:
                rows.append([c.text.strip() for c in row.cells])
            if not rows:
                continue
            headers = rows[0]
            data_rows = rows[1:] if len(rows) > 1 else []
            tables.append(
                ParsedTable(
                    headers=headers,
                    rows=data_rows,
                )
            )

        # docx 没有页概念；用单页代表整篇
        single_page = ParsedPage(
            page_no=1,
            text_blocks=text_blocks,
            tables=tables,
        )
        populate_l0_blocks(single_page, doc_id=ctx.doc_id, parser=self.name)

        return ParsedDoc(
            doc_id=ctx.doc_id,
            source_path=str(path),
            pages=[single_page],
            metadata=ctx.metadata,
            toc=toc,
            parser=self.name,
            parser_version=self.version,
        )

    # ---------- 工具 ----------

    @staticmethod
    def _heading_level(style_name: str) -> int | None:
        """从 style 名识别标题层级。

        python-docx 的 style 名通常是 'Heading 1' / 'Heading 2' / 'Title' 等。
        """
        s = style_name.lower().strip()
        if s.startswith("heading "):
            try:
                lvl = int(s.split(" ", 1)[1].strip())
                if 1 <= lvl <= 9:
                    return lvl
            except (ValueError, IndexError):
                return None
        if s in ("title", "subtitle"):
            return 1
        return None
